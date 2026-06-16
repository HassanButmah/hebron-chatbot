"""Insert student quick-start guide after section 5.1 and fix section numbering."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    para = Paragraph(new_p, paragraph._parent)
    if text:
        para.add_run(text)
    if style:
        para.style = style
    return para


def insert_table_after(doc: Document, paragraph: Paragraph, headers: list[str], rows: list[list[str]]) -> Paragraph:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        for cell in table.rows[0].cells[i].paragraphs:
            for run in cell.runs:
                run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = value
    paragraph._element.addnext(table._tbl)
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_caption_after(anchor: Paragraph, text: str) -> Paragraph:
    para = insert_paragraph_after(anchor, text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.italic = True
        run.font.size = Pt(10)
    return para


def add_picture_after(anchor: Paragraph, image_path: Path) -> Paragraph:
    para = insert_paragraph_after(anchor, "")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image_path), width=Inches(5.8))
    return para


def find_heading(doc: Document, prefix: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip().startswith(prefix):
            return p
    raise ValueError(f"Heading not found: {prefix}")


def renumber_headings(doc: Document) -> None:
    replacements = [
        ("5.6 التحديات", "5.8 التحديات"),
        ("5.5 البنية التحتية", "5.7 البنية التحتية"),
        ("5.5.1 لغات", "5.6.1 لغات"),
        ("5.4.7 الذكاء", "5.6.7 الذكاء"),
        ("5.4.6 قواعد", "5.6.6 قواعد"),
        ("5.4.5 المكتبات", "5.6.5 المكتبات"),
        ("5.4.4 أدوات إدارة", "5.6.4 أدوات إدارة"),
        ("5.4.3 أدوات التصميم", "5.6.3 أدوات التصميم"),
        ("5.4.2 أدوات التحرير", "5.6.2 أدوات التحرير"),
        ("5.5 التطوير", "5.6 التطوير"),
        ("5.4  لماذا", "5.5  لماذا"),
        ("5.4 لماذا", "5.5 لماذا"),
        ("5.3 آلية", "5.4 آلية"),
        ("5.2 المسار", "5.3 المسار"),
    ]
    for p in doc.paragraphs:
        if not p.style.name.startswith("Heading"):
            continue
        for old, new in replacements:
            if old in p.text:
                p.text = p.text.replace(old, new, 1)
                break


def renumber_legacy_figures(doc: Document) -> None:
    replacements = [
        ("جدول رقم (5.1)", "جدول رقم (5.3)"),
        ("جدول (5.1): المكتبات", "جدول (5.3): المكتبات"),
        ("الشكل رقم (5.1)", "الشكل رقم (5.2)"),
        ("الشكل (5.1): مراحل التنفيذ", "الشكل (5.2): مراحل التنفيذ"),
        ("المخطط رقم (5.6)", "المخطط رقم (5.2)"),
    ]
    for p in doc.paragraphs:
        for old, new in replacements:
            if old in p.text:
                p.text = p.text.replace(old, new)


def insert_quick_start(doc: Document, folder: Path) -> None:
    # Insert after the last paragraph of section 5.1 (before old 5.2 المسار)
    anchor = find_heading(doc, "5.2 المسار")
    prev = doc.paragraphs[[p._element for p in doc.paragraphs].index(anchor._element) - 1]

    image_path = folder / "Capture.PNG"
    if not image_path.exists():
        image_path = folder / "Capture.png"

    cur = prev

    def h1(text: str) -> None:
        nonlocal cur
        cur = insert_paragraph_after(cur, text, "Heading 1")

    def body(text: str) -> None:
        nonlocal cur
        cur = insert_paragraph_after(cur, text, "Normal")

    def caption(text: str) -> None:
        nonlocal cur
        cur = add_caption_after(cur, text)

    def table(headers: list[str], rows: list[list[str]]) -> None:
        nonlocal cur
        cur = insert_table_after(doc, cur, headers, rows)

    def picture(path: Path) -> None:
        nonlocal cur
        cur = add_picture_after(cur, path)

    h1("5.2 دليل البدء السريع للطلاب")
    body(
        "يهدف هذا القسم إلى تمكين الطلاب والباحثين في الجامعة من تشغيل المشروع من الصفر "
        "واستخدامه كمرجع عملي. للتفاصيل الكاملة راجع ملف README.md في جذر المشروع."
    )

    h1("5.2.1 المتطلبات الأساسية")
    body("قبل البدء، يجب تثبيت البرامج التالية على الجهاز أو الخادم المستخدم:")
    table(
        ["الأداة", "الإصدار", "الغرض"],
        [
            ["Python", "3.10+", "Backend ومسار RAG"],
            ["PostgreSQL", "14+", "قاعدة البيانات العلائقية"],
            ["Ollama", "latest", "التضمين (bge-m3) ونموذج LLM محلي"],
            ["Node.js", "18+", "لوحة الإدارة (React)"],
            ["Git", "latest", "إدارة نسخ المشروع"],
        ],
    )

    h1("5.2.2 استنساخ المشروع وإعداد البيئة")
    body(
        "1. انسخ مجلد المشروع (أو استنسخه من مستودع GitHub الخاص بالفريق).\n"
        "2. أنشئ بيئة Python:\n"
        "   conda create -n arabic-rag python=3.10 -y\n"
        "   conda activate arabic-rag\n"
        "3. ثبّت المكتبات:\n"
        "   pip install -r requirements.txt\n"
        "4. انسخ ملف الإعدادات:\n"
        "   copy scripts\\.env.example .env   (Windows)\n"
        "   cp scripts/.env.example .env      (Linux/macOS)\n"
        "5. عدّل ملف .env وأدخل بيانات الاتصال بقاعدة PostgreSQL وباقي الإعدادات."
    )

    h1("5.2.3 البرامج والتقنيات المستخدمة")
    body(
        "يوضح الشكل (5.1) المكوّنات التقنية الرئيسية للمشروع. يربط الجدول (5.1) كل مكوّن "
        "بموقعه داخل المشروع وخطوة الإعداد الأساسية."
    )
    caption("شكل (5.1): البرامج والتقنيات والخوارزميات المستخدمة في المشروع")
    picture(image_path)
    caption("جدول (5.1): ربط المكوّنات التقنية بملفات المشروع وخطوات الإعداد")
    table(
        ["المكوّن", "الموقع في المشروع", "خطوة الإعداد"],
        [
            ["Python", "rag_api.py ، src/", "pip install -r requirements.txt"],
            ["Node.js", "admin-panel/", "npm install"],
            ["React / HTML / CSS / JS", "admin-panel/ ، widget/", "npm run dev"],
            ["Flask", "rag_api.py", "python rag_api.py"],
            ["PostgreSQL", "database.py", "DATABASE_URL في .env"],
            ["ChromaDB", "chroma_db/", "تُنشأ تلقائياً عند التشغيل"],
            ["Ollama + BGE-M3", "rag_system.py", "ollama pull bge-m3"],
            ["DeepSeek API", ".env", "LLM_PROVIDER=openai_compatible + LLM_API_KEY"],
            ["Ollama (LLM محلي)", ".env", "LLM_PROVIDER=ollama"],
            ["GitHub", "المستودع", "git clone أو نسخ المجلد"],
            ["VS Code", "—", "فتح مجلد المشروع"],
        ],
    )

    h1("5.2.4 ربط الأدوات بملف الإعدادات (.env)")
    body("يوضح الجدول (5.2) أهم المتغيرات في ملف .env وكيفية ربط كل أداة بالنظام:")
    caption("جدول (5.2): أهم متغيرات ملف .env")
    table(
        ["المتغير", "الأداة", "مثال / قيمة"],
        [
            ["DATABASE_URL", "PostgreSQL", "postgresql://user:pass@localhost:5432/hebron_rag"],
            ["EMBED_MODEL", "Ollama", "bge-m3"],
            ["OLLAMA_BASE_URL", "Ollama", "http://localhost:11434"],
            ["LLM_PROVIDER", "نموذج اللغة", "ollama أو openai_compatible"],
            ["LLM_API_KEY", "DeepSeek / OpenAI", "مفتاح API من لوحة المزود"],
            ["LLM_API_BASE_URL", "DeepSeek", "https://api.deepseek.com"],
            ["LLM_MODEL", "DeepSeek", "deepseek-chat أو deepseek-v4-flash"],
            ["JWT_SECRET_KEY", "لوحة الإدارة", "سلسلة سرية عشوائية"],
            ["REDIS_URL", "Rate limiting", "fakeredis (تطوير) أو redis://localhost:6379/0"],
        ],
    )

    h1("5.2.5 ترتيب تشغيل النظام")
    body(
        "1. تشغيل خدمة PostgreSQL والتأكد من صحة DATABASE_URL.\n"
        "2. تشغيل Ollama وتنزيل نموذج التضمين: ollama pull bge-m3\n"
        "3. تشغيل الخادم الرئيسي: python rag_api.py  (المنفذ 5000)\n"
        "4. (اختياري) تشغيل Mock API للبيانات الحية: python mock_api_server.py  (المنفذ 5001)\n"
        "5. تشغيل لوحة الإدارة: cd admin-panel ثم npm run dev  (المنفذ 5173)\n"
        "6. فتح المتصفح على http://localhost:5173 وتسجيل الدخول."
    )

    h1("5.2.6 أول اختبار للنظام")
    body(
        "بعد التشغيل، يُنصح بتنفيذ الخطوات التالية للتحقق من عمل المشروع:\n"
        "• تسجيل الدخول إلى لوحة الإدارة (المستخدم الافتراضي: ChatBot / Hebron@uni — "
        "يُفضّل تغيير كلمة المرور فوراً).\n"
        "• رفع ملف PDF أو Word من قسم إدارة الملفات.\n"
        "• طرح سؤال من واجهة المحادثة (Widget) أو من صفحة الاختبار.\n"
        "• التحقق من أن الإجابة مبنية على محتوى الملف المرفوع."
    )
    body(
        "ملاحظة: عند استخدام DeepSeek API، يُرسل كل سؤال مستخدم عبر استدعاءين للـ API: "
        "أولاً لإعادة صياغة السؤال (Query Rewrite)، ثم لتوليد الإجابة عبر مسار RAG."
    )


def main() -> None:
    folder = Path(__file__).resolve().parent
    docx_files = [f for f in folder.glob("*.docx") if "backup" not in f.name]
    if not docx_files:
        raise SystemExit("No .docx file found")

    source = docx_files[0]
    backup = source.with_suffix(".backup.docx")
    shutil.copy2(backup, source)

    doc = Document(source)
    insert_quick_start(doc, folder)
    renumber_headings(doc)
    renumber_legacy_figures(doc)
    doc.save(source)

    print(f"Updated: {source}")


if __name__ == "__main__":
    main()
