"""Generate Word document for the LLM token usage and cost study (testing chapter)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = value
    doc.add_paragraph()


def main() -> None:
    out_path = Path(__file__).resolve().parent / "LLM_Token_Usage_Cost_Study_May2026.docx"

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("LLM Token Usage and Cost Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Hebron University Chatbot — Testing Phase (May 2026)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Period: May 2026\n").bold = False
    meta.add_run("Primary model: DeepSeek V4 Flash\n")
    meta.add_run("Data source: DeepSeek API usage exports (amount-2026-5.csv, cost-2026-5.csv)")

    doc.add_paragraph()

    # Section 1
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "During system testing, the chatbot was connected to the DeepSeek API to evaluate "
        "token consumption and operational cost per user interaction. Each user message triggers "
        "two API calls:"
    )
    doc.add_paragraph(
        "Query rewrite — reformulates the user's question using conversation history "
        "(approximately 500 input tokens, 25 output tokens).",
        style="List Number",
    )
    doc.add_paragraph(
        "RAG answer generation — retrieves context and produces the final reply "
        "(approximately 2,000 input tokens, 85 output tokens).",
        style="List Number",
    )
    doc.add_paragraph(
        "This section reports measured usage from May 2026 and compares the cost if the same "
        "traffic were processed by alternative models."
    )

    # Section 2
    doc.add_heading("2. Measured Usage (DeepSeek V4 Flash)", level=1)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Total API requests", "868"],
            ["Estimated user messages (868 ÷ 2)", "434"],
            ["Active testing days", "15"],
            ["Total tokens processed", "1,142,115"],
            ["Total cost", "$0.12"],
        ],
    )

    doc.add_heading("Per user message (1 question + 1 answer)", level=2)
    add_table(
        doc,
        ["Metric", "Measured average"],
        [
            ["Input tokens (cache miss)", "~1,797"],
            ["Input tokens (cache hit)", "~725"],
            ["Output tokens", "~110"],
            ["Total billable tokens", "~2,632"],
            ["Cost per message", "$0.00028"],
            ["Cost per 1,000 messages", "$0.28"],
        ],
    )

    doc.add_paragraph(
        "The rewrite call uses significantly fewer tokens than the RAG call. When both calls are "
        "combined, each user interaction consumes approximately 2,632 billable tokens in total."
    )

    # Section 3
    doc.add_heading("3. Model Cost Comparison", level=1)
    doc.add_paragraph(
        "The table below uses the measured token profile per user message and published API rates "
        "(June 2026)."
    )
    add_table(
        doc,
        ["Model", "Cost per message", "Cost per 1,000 messages", "Relative to Flash"],
        [
            ["DeepSeek V4 Flash", "$0.00028", "$0.28", "1.0×"],
            ["OpenAI GPT-4o mini", "$0.00039", "$0.39", "1.4×"],
            ["DeepSeek V4 Pro", "$0.00088", "$0.88", "3.1×"],
            ["OpenAI GPT-4o", "$0.00650", "$6.50", "23×"],
        ],
    )

    p = doc.add_paragraph()
    p.add_run("Pricing sources: ").italic = True
    p.add_run("DeepSeek API (api-docs.deepseek.com), OpenAI API (openai.com/api/pricing). ")
    p.add_run(
        "OpenAI figures are projected at the same token volume; actual Arabic tokenization may vary slightly."
    ).italic = True

    doc.add_heading("Projected monthly cost (DeepSeek V4 Flash)", level=2)
    add_table(
        doc,
        ["User messages / month", "Estimated LLM cost"],
        [
            ["1,000", "$0.28"],
            ["10,000", "$2.84"],
            ["50,000", "$14.21"],
            ["100,000", "$28.43"],
        ],
    )

    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run("LLM API cost only; excludes hosting, embeddings, and infrastructure.")

    # Section 4
    doc.add_heading("4. DeepSeek V4 Pro (Evaluation Sample)", level=1)
    doc.add_paragraph(
        "A small Pro sample (61 API requests, $0.05) was used for model comparison during testing. "
        "Pro averaged 329 output tokens per request versus 55 for Flash, and cost approximately "
        "5× more per API call. It is suitable for evaluation but not recommended as the default "
        "production model."
    )

    # Section 5
    doc.add_heading("5. Findings", level=1)
    findings = [
        "Each user message consumes approximately 2,632 billable tokens across two API calls "
        "(rewrite + RAG).",
        "At measured rates, DeepSeek V4 Flash costs approximately $0.28 per 1,000 messages during testing.",
        "DeepSeek V4 Flash remains the most cost-effective option for production deployment.",
        "GPT-4o mini would cost approximately 1.4× more; DeepSeek V4 Pro approximately 3.1× more; "
        "GPT-4o approximately 23× more at the same token volume.",
    ]
    for item in findings:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run("Note: ").bold = True
    closing.add_run(
        "May 2026 data reflects testing traffic over 15 active days (434 user messages). "
        "Production volumes may differ, but the per-message token and cost profile provides a "
        "reliable baseline for budget planning."
    )

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
